"""
Resume document I/O helpers (``.docx`` only, via python-docx).

- ``extract_text`` reads an uploaded resume into plain text for the agents.
- ``build_docx`` renders a tailored-resume dict (as produced by
  ``AIJobAgent.tailor_resume``) back into a clean ``.docx`` file.

Legacy binary ``.doc`` files are not supported — convert them to ``.docx`` first.
"""

from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt


def extract_text(data: bytes) -> str:
    """Extract plain text from ``.docx`` bytes (paragraphs and table cells)."""
    document = Document(BytesIO(data))

    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    # Include table content (many resumes use tables for layout).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))

    return '\n'.join(lines)


def build_docx(resume: Dict[str, Any]) -> bytes:
    """Render a tailored-resume dict into a ``.docx`` file, returned as bytes.

    Expected shape (all fields optional)::

        {
          "name": str,
          "contact": str,
          "summary": str,
          "sections": [{"heading": str, "bullets": [str, ...]}, ...]
        }
    """
    document = Document()

    name = (resume.get('name') or '').strip()
    if name:
        document.add_heading(name, level=0)

    contact = (resume.get('contact') or '').strip()
    if contact:
        contact_paragraph = document.add_paragraph()
        run = contact_paragraph.add_run(contact)
        run.font.size = Pt(10)

    summary = (resume.get('summary') or '').strip()
    if summary:
        document.add_heading('Summary', level=1)
        document.add_paragraph(summary)

    for section in resume.get('sections') or []:
        if not isinstance(section, dict):
            continue
        heading = (section.get('heading') or '').strip()
        if heading:
            document.add_heading(heading, level=1)
        for bullet in section.get('bullets') or []:
            text = str(bullet).strip()
            if text:
                document.add_paragraph(text, style='List Bullet')

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
